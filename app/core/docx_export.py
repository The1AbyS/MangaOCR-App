from pathlib import Path

from docx import Document


def export_ocr_cache_to_docx(entries, ocr_cache, filename: str):
    doc = Document()

    for path in entries:
        cached = ocr_cache.get_for_path(path)
        if not cached:
            continue

        boxes, frames, md5 = cached
        text = "\n".join([box.text for box in boxes if box.text])
        doc.add_paragraph(f"{Path(path).name}")
        doc.add_paragraph(text)
        doc.add_paragraph("")

    doc.save(filename)
