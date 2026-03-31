import re
import os
from pathlib import Path
from docx import Document
from PySide6.QtWidgets import (QMainWindow, QWidget, QListWidget, QVBoxLayout, QLabel,
                               QToolBar, QFileDialog, QStatusBar, QGraphicsScene, QMenu,
                               QInputDialog, QApplication, QDockWidget, QToolButton, QStackedWidget, QListWidgetItem)
from PySide6.QtGui import QPixmap, QAction, QPainter, QPen, QColor, QDesktopServices, QFont, QPainterPath, QIcon
from PySide6.QtCore import Qt, QPoint, QRectF, QSettings, QSize, QUrl
from PIL import Image
from qt_material_icons import MaterialIcon
from .jardic import JardicWidget
from .preview import ImageView
from .textexportpanel import TextExportPanel
from .yolosettings import YoloSettingsWidget
from .projectmanager import ProjectManager
from ..core.cahcefolder import CacheFolder
from ..core.cache import OCRCache
from ..core.utils import natural_key
from ..core.threads import ModelsLoadThread
from ..core.ocr import OCRThread, BatchThread
from ..core.parser import ImageParser
from ..ignore import ignore_warnings
from ..core.rawkuma import SearchWindowRawkuma

ignore_warnings()

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".webp"}

class MainWindow(QMainWindow):
    __version__ = "Alpha 0.2.3.0"

    def __init__(self):
        super().__init__()
        self.setWindowTitle("MangaOCR App")
        self.settings = QSettings("MangaOCR App")

        self.setDockOptions(QMainWindow.AnimatedDocks)
        self.setDockNestingEnabled(True)

        self._create_widgets()
        self._create_central()
        self._create_statusbar()
        self._create_actions()
        self._create_toolbar()
        try:
            self.restore_window_state()
        except Exception as e:
            print("Ошибка восстановления состояния окна:", e)
            pass

        self.start_models_load()
        self.cache_folder = CacheFolder()

        self.mocr = None
        self.yolo_detector = None

        self.ocr_cache = OCRCache()
        self.ocr_cache._load_cache()

        self.current_preview_idx = None
        self.current_folder = None
        self.entries = [] 

    def restore_window_state(self):
        size = self.settings.value("window_size", QSize(1000, 700), type=QSize)
        pos  = self.settings.value("window_pos", QPoint(100, 100), type=QPoint)

        self.resize(size)
        self.move(pos)

        was_maximized = self.settings.value("window_maximized", False, type=bool)
        if was_maximized:
            self.setWindowState(self.windowState() | Qt.WindowMaximized)

        s = QSettings("MangaOCR", "Layout")
        state = s.value("state")
        if state:
            self.restoreState(state)

    def closeEvent(self, event):
        self.settings.setValue("window_maximized", self.isMaximized())

        if not self.isMaximized():
            self.settings.setValue("window_size", self.size())
            self.settings.setValue("window_pos", self.pos())

        s = QSettings("MangaOCR", "Layout")
        s.setValue("state", self.saveState())

        super().closeEvent(event)

    def start_models_load(self):
        self.models_thread = ModelsLoadThread()
        self.models_thread.finished.connect(self.on_models_loaded)
        self.models_thread.error.connect(lambda e: print("Ошибка загрузки моделей:", e))
        self.models_thread.start()

    def on_models_loaded(self, mocr, yolo):
        self.mocr = mocr
        self.yolo_detector = yolo
        self.statusBar().showMessage("Модели успешно загружены")

    def _create_actions(self):
        self.open_folder_act = QAction("Открыть папку", self)
        self.open_folder_act.setShortcut("Ctrl+O")
        self.open_folder_act.triggered.connect(self.action_open_folder)

        self.recent_menu = QMenu(self)
        self.recent_menu.aboutToShow.connect(self.update_recent_menu)
        self.open_folder_act.setMenu(self.recent_menu)

        self.show_frames_act = QAction("Отображение", self)

        self.show_frames_menu = QMenu(self)
        self.show_frames_toggle = QAction("Фреймы", self, checkable=True)
        self.show_frames_toggle.setChecked(False)
        self.show_frames_toggle.triggered.connect(self.toggle_show_frames)

        self.show_numbers_toggle = QAction("Номера боксов", self, checkable=True)
        self.show_numbers_toggle.setChecked(False)
        self.show_numbers_toggle.triggered.connect(self.toggle_show_numbers_boxes)

        self.show_frames_menu.addAction(self.show_frames_toggle)
        self.show_frames_menu.addAction(self.show_numbers_toggle)
        self.show_frames_act.setMenu(self.show_frames_menu)

        self.batch_act = QAction("Обработать всё", self)
        self.batch_act.triggered.connect(self.action_batch_process)

        self.export_act = QAction("Сохранить как...", self)
        self.export_act.triggered.connect(self.action_export_text)

        self.previous_image_act = QAction("Предыдущее изображение", self)
        self.previous_image_act.setShortcut("Up")
        self.previous_image_act.triggered.connect(self.previous_image)
        self.addAction(self.previous_image_act)

        self.next_image_act = QAction("Следующее изображение", self)
        self.next_image_act.setShortcut("Down")
        self.next_image_act.triggered.connect(self.next_image)
        self.addAction(self.next_image_act)

        self.parser_act = QAction("Загрузить изображения из URL", self)
        self.parser_act.triggered.connect(self.action_parser)

        self.clipboard_act = QAction("Буфер обмена", self)
        self.clipboard_act.triggered.connect(self.clipboard_action)
        self.addAction(self.clipboard_act)
        self.clipboard_act.setShortcut("Ctrl+V")

        self.refresh_current_act = QAction("Сбросить кэш", self)
        self.refresh_current_act.triggered.connect(self.refresh_current)
        self.addAction(self.refresh_current_act)
        self.refresh_current_act.setShortcut("Ctrl+R")

        self.delete_current_act = QAction("Удалить текущую страницу из списка", self)
        self.delete_current_act.triggered.connect(self.delete_from_list)
        self.addAction(self.delete_current_act)
        self.delete_current_act.setShortcut("Del")

        self.window_button = QToolButton()
        self.window_button.setText("Окна")
        self.window_button.setPopupMode(QToolButton.MenuButtonPopup)
        self.window_button.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        self.view_menu = QMenu()
        self.view_menu.addAction(self.dock_list.toggleViewAction())
        self.view_menu.addAction(self.dock_text.toggleViewAction())
        self.view_menu.addAction(self.dock_dict.toggleViewAction())
        self.view_menu.addAction(self.dock_rawkuma.toggleViewAction())
        self.view_menu.addAction(self.dock_yolo_settings.toggleViewAction())
        self.window_button.setMenu(self.view_menu)

    def _create_toolbar(self):
        tb = QToolBar("Main")
        tb.setObjectName("main_toolbar")
        tb.setMovable(False)

        self.addToolBar(tb)

        tb.addAction(self.open_folder_act)
        tb.addAction(self.show_frames_act)
        tb.addAction(self.batch_act)
        tb.addAction(self.export_act)
        tb.addWidget(self.window_button)

    def _create_widgets(self):

        self.list_widget = QListWidget()
        self.list_widget.itemClicked.connect(self.on_item_clicked)
        self.list_widget.setMinimumWidth(300)
        self.list_widget.setContextMenuPolicy(Qt.CustomContextMenu)
        self.list_widget.customContextMenuRequested.connect(self.show_list_menu)

        self.project_manager = ProjectManager()
        self.project_manager.projectOpened.connect(self.load_folder)

        self.text_export_panel = TextExportPanel()

        self.jardic_widget = JardicWidget(self)

        self.scene = None
        self.preview_view = ImageView()

        self.rawkuma_widget = SearchWindowRawkuma()
        self.rawkuma_widget.chapter_clicked.connect(self.action_parser_from_rawkuma)

        self.yolo_settings = YoloSettingsWidget()
        self.yolo_settings.settings_changed.connect(self.update_yolo_params)

    def action_parser_from_rawkuma(self, url):
        self.action_parser(text=url)

    def update_yolo_params(self, settings):
        self.current_yolo_params = settings
        try:
            self.refresh_current()
        except Exception as e:
            self.statusBar().showMessage(f"Ошибка обновления YOLO: {e}")

    def _create_central(self):

        self.preview_container = QWidget()
        self.preview_layout = QVBoxLayout(self.preview_container)
        self.preview_layout.setContentsMargins(4, 4, 4, 4)

        self.scene = QGraphicsScene()
        self.preview_view.setScene(self.scene)

        self.preview_layout.addWidget(self.preview_view)

        self.central_stack = QStackedWidget()
        self.central_stack.addWidget(self.project_manager)  
        self.central_stack.addWidget(self.preview_container)

        self.setCentralWidget(self.central_stack)

        self.central_stack.setCurrentIndex(0)


        self.dock_list = QDockWidget("Страницы", self)
        self.dock_list.setWidget(self.list_widget)
        self.dock_list.setFeatures(
            QDockWidget.DockWidgetMovable |
            QDockWidget.DockWidgetClosable
        )
        self.dock_list.setObjectName("dock_pages")

        self.addDockWidget(Qt.LeftDockWidgetArea, self.dock_list)

        self.dock_text = QDockWidget("Текст", self)
        self.dock_text.setWidget(self.text_export_panel)
        self.dock_text.setFeatures(
            QDockWidget.DockWidgetMovable |
            QDockWidget.DockWidgetClosable
        )
        self.dock_text.setObjectName("dock_text")

        self.addDockWidget(Qt.RightDockWidgetArea, self.dock_text)

        self.dock_dict = QDockWidget("Jardic", self)
        self.dock_dict.setWidget(self.jardic_widget)
        self.dock_dict.setFeatures(
            QDockWidget.DockWidgetMovable |
            QDockWidget.DockWidgetClosable
        )
        self.dock_dict.setObjectName("dock_dict")

        self.addDockWidget(Qt.RightDockWidgetArea, self.dock_dict)

        self.dock_rawkuma = QDockWidget("Rawkuma", self)
        self.dock_rawkuma.setWidget(self.rawkuma_widget)
        self.dock_rawkuma.setFeatures(
            QDockWidget.DockWidgetMovable |
            QDockWidget.DockWidgetClosable
        )
        self.dock_rawkuma.setObjectName("dock_rawkuma")

        self.addDockWidget(Qt.RightDockWidgetArea, self.dock_rawkuma)

        self.dock_yolo_settings = QDockWidget("Настройки YOLO", self)
        self.dock_yolo_settings.setWidget(self.yolo_settings)
        self.dock_yolo_settings.setFeatures(
            QDockWidget.DockWidgetMovable |
            QDockWidget.DockWidgetClosable
        )
        self.dock_yolo_settings.setObjectName("dock_yolo_settings")

        self.addDockWidget(Qt.RightDockWidgetArea, self.dock_yolo_settings)

    def _create_statusbar(self):
        sb = QStatusBar()
        self.setStatusBar(sb)
        status = self.statusBar()

        self.version_label = QLabel(self.__version__)
        self.version_label.setStyleSheet("color: gray; padding-right: 10px;")
        sb.addPermanentWidget(self.version_label)

    def action_open_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Выберите папку с изображениями")
        if folder: 
            self.cache_folder.add(Path(folder))
            self.load_folder(Path(folder))

    def clipboard_action(self):
        clipboard = QApplication.clipboard()
        text = clipboard.text().strip()

        if not text:
            self.statusBar().showMessage("Буфер обмена пуст")
            return

        if text.startswith("http://") or text.startswith("https://") or text.startswith("www."):
            self.statusBar().showMessage("Парсинг URL из буфера обмена...")
            try:
                self.action_parser(text=text)
            except Exception as e:
                self.statusBar().showMessage(f"Ошибка парсинга URL: {e}")
            return

        text = re.sub(r"^file:/+", "", text)
        path_clipboard = Path(text)

        if path_clipboard.exists():
            try:
                if path_clipboard.is_dir():
                    self.cache_folder.add(path_clipboard)
                    self.load_folder(path_clipboard)
            except Exception as e:
                self.statusBar().showMessage(f"Ошибка добавления папки: {e}")
        else:
            self.statusBar().showMessage("Буфер обмена не является ни URL, ни путём")

    def action_parser(self, bool=None, text=None):
        if text is None:
            url, ok = QInputDialog.getText(self, "Загрузить изображения из URL", "Введите URL страницы:")
            if not ok or not url.strip():
                return
        else:
            url = text.strip()

        self.thread = ImageParser(url)
        self.thread.progress.connect(self.statusBar().showMessage)
        self.thread.finished.connect(self.on_images_downloaded)
        self.thread.start()

    def on_images_downloaded(self, saved_files, out_dir):
        if saved_files and out_dir:
            self.cache_folder.add(Path(out_dir))
            self.load_folder(Path(out_dir))

    def toggle_show_frames(self, checked):
        self.show_frames = checked
        self.show_frames_and_numbers_boxes()

    def toggle_show_frames(self, checked):
        self.show_frames = checked
        self.show_frames_and_numbers_boxes()

    def toggle_show_numbers_boxes(self, checked):
        self.show_numbers_boxes = checked
        self.show_frames_and_numbers_boxes()

    def show_frames_and_numbers_boxes(self):
        current_index = self.list_widget.currentRow()

        if current_index < 0:
            return

        frames = self.frames if getattr(self, 'show_frames', False) else None
        numbers_boxes = self.text_boxes if getattr(self, 'show_numbers_boxes', False) else None

        self.show_preview(
            self.entries[current_index],
            boxes=getattr(self, 'text_boxes', None),
            frames=frames,
            numbers_boxes=numbers_boxes
        )

    def update_recent_menu(self):
        self.recent_menu.clear()
        folders = self.cache_folder.get_all()
        if not folders:
            self.recent_menu.addAction("(пусто)").setEnabled(False)
        else:
            for path in folders:
                act = self.recent_menu.addAction(str(path))
                act.triggered.connect(lambda checked=False, p=path: self.load_folder(p))
            self.recent_menu.addSeparator()

    def load_folder(self, folder: Path):
        try:
            folder = Path(folder) 
            if not folder or not folder.exists() or not folder.is_dir():
                self.statusBar().showMessage("Папка не существует")
                self.list_widget.clear()
                self.entries = []
                self.current_folder = None
                return

            self.cache_folder.add(folder)
            self.current_folder = folder
        except Exception as e:
            self.statusBar().showMessage(f"Ошибка загрузки папки: {e}")
            self.list_widget.clear()
            self.entries = []
            self.current_folder = None
            return

        self.entries = [
            p for p in sorted(folder.iterdir(), key=lambda x: natural_key(x.name))
            if p.is_file() and p.suffix.lower() in IMAGE_EXTENSIONS
        ]

        self.list_widget.clear()
        self.dock_list.show()
        for p in self.entries:
            item = QListWidgetItem(p.name)
            item.setData(Qt.UserRole, str(p)) 
            self.list_widget.addItem(item)

        if self.entries:
            self.list_widget.setCurrentRow(0)

        self.statusBar().showMessage(f"Загружено {len(self.entries)} изображений из {folder} (текущий уровень)")

    def on_item_clicked(self, item):
        if self.yolo_detector is None:
            self.statusBar().showMessage("Модели ещё не загружены. Пожалуйста, подождите.")
            return

        idx = self.list_widget.row(item)
        if idx < 0 or idx >= len(self.entries):
            return

        path = self.entries[idx]
        self.show_preview(path, boxes=None, reset_zoom=True) 

        self.central_stack.setCurrentIndex(1)

        self._current_image_token = object()
        self.is_cached(path = path, idx = idx)

        if hasattr(self, 'ocr_thread') and self.ocr_thread is not None:
            self.ocr_thread.quit()
            self.ocr_thread.wait()

        self.ocr_thread = OCRThread(self, path, token=self._current_image_token)
        self.ocr_thread.finished.connect(self.on_ocr_finished)
        self.ocr_thread.start()

        self.current_preview_idx = idx

    def is_cached(self, path, idx):
        if hasattr(self, 'ocr_cache'):
            try:
                cached = self.ocr_cache.get_for_path(path)
            except Exception:
                cached = None
            if cached:
                boxes, frames, md5 = cached
                self.text_boxes, self.frames = boxes, frames
                self.show_preview(path, boxes=self.text_boxes, frames=self.frames, reset_zoom=False)
                self.text_export_panel.set_boxes(self.text_boxes, frames=self.frames, path=self.entries[idx])
                return

    def next_image(self):
        current_index = self.list_widget.currentRow()
        if current_index < self.list_widget.count() - 1:
            self.list_widget.setCurrentRow(current_index + 1)
            self.on_item_clicked(self.list_widget.currentItem())

    def previous_image(self):
        current_index = self.list_widget.currentRow()
        if current_index > 0:
            self.list_widget.setCurrentRow(current_index - 1)
            self.on_item_clicked(self.list_widget.currentItem())

    def show_list_menu(self, pos):
        item = self.list_widget.itemAt(pos)
        if item is None:
            return

        idx = self.list_widget.row(item)
        if idx < 0 or idx >= self.list_widget.count():
            return

        path = Path(item.data(Qt.UserRole))

        menu = QMenu()
        open_action = menu.addAction("Открыть как файл")
        open_action.triggered.connect(lambda: QDesktopServices.openUrl(QUrl.fromLocalFile(str(path))))
        open_icon = MaterialIcon("image")
        open_action.setIcon(open_icon)

        open_folder_action = menu.addAction("Открыть папку")
        open_folder_action.triggered.connect(lambda: QDesktopServices.openUrl(QUrl.fromLocalFile(str(path.parent))))
        folder_icon = MaterialIcon("folder")
        open_folder_action.setIcon(folder_icon)

        refresh_action = menu.addAction("Сбросить кэш")
        refresh_action.triggered.connect(lambda: self.refresh_current(path=path))
        refresh_icon = MaterialIcon("refresh")
        refresh_action.setIcon(refresh_icon)

        delete_action = menu.addAction("Удалить из списка")
        delete_action.triggered.connect(lambda: self.delete_from_list(idx))
        delete_icon = MaterialIcon("delete")
        delete_action.setIcon(delete_icon)

        menu.addSeparator()

        info_action = menu.addAction("Информация о файле")
        info_action.triggered.connect(lambda: self.show_file_info(path))
        info_icon = MaterialIcon("info")
        info_action.setIcon(info_icon)

        menu.exec(self.list_widget.viewport().mapToGlobal(pos))

    def action_batch_process(self):
        self.batch_thread = BatchThread(self, self.entries, parent=self)

        self.batch_thread.item_started.connect(self._on_batch_item_started)
        self.batch_thread.item_finished.connect(self._on_batch_item_finished)
        self.batch_thread.all_done.connect(self._on_batch_done)

        self.batch_thread.start()

    def _on_batch_item_started(self, idx, path):
        self.list_widget.setCurrentRow(idx)
        self.statusBar().showMessage(f"Обрабатывается: {Path(self.list_widget.item(idx).data(Qt.UserRole)).name} ({idx+1}/{self.list_widget.count()})")

    def _on_batch_item_finished(self, idx, result):
        boxes, frames = result
        self.text_boxes, self.frames = boxes, frames
        self.text_export_panel.set_boxes(boxes, frames=frames)
        self.show_preview(Path(self.list_widget.item(idx).data(Qt.UserRole)), boxes=self.text_export_panel._boxes, frames=frames, reset_zoom=False)
        self.statusBar().showMessage(f"Завершено: {Path(self.list_widget.item(idx).data(Qt.UserRole)).name} ({idx+1}/{self.list_widget.count()})")

    def _on_batch_done(self):
        self.statusBar().showMessage("Пакетная обработка завершена")

    def delete_from_list(self, bool=None, idx=None):
        if idx is None:
            idx = self.list_widget.currentRow()
        if idx < 0 or idx >= self.list_widget.count():
            return

        self.list_widget.takeItem(idx)
        self.entries.pop(idx)

        if idx == getattr(self, 'current_preview_idx', None) or self.list_widget.count() == 0:
            self.show_preview(QPixmap())
            self.statusBar().showMessage("")
            self.text_export_panel.set_boxes([], [])

    def show_file_info(self, path):
        if not path.exists():
            self.statusBar().showMessage("Файл не найден")
            return

        def readable_size(size_bytes: int) -> str:
            if size_bytes == 0:
                return "0 байт"
            units = ["байт", "КБ", "МБ", "ГБ", "ТБ"]
            i = 0
            size = float(size_bytes)
            while size >= 1024 and i < len(units) - 1:
                size /= 1024
                i += 1
            return f"{size:.1f} {units[i]}" if i > 0 else f"{int(size)} {units[i]}"

        def get_file_info(path: Path) -> str:
            if not path.is_file():
                return "Файл не найден"

            stat = path.stat()
            size_str = readable_size(stat.st_size)

            with Image.open(path) as img:
                    w, h = img.size

            lines = [
                f"{'Имя:'} {path.name}",
                f"{'Размер:'} {size_str} ({stat.st_size:,} байт)",
                f"{'Путь:'} {path.resolve()}",
                f"{'Разрешение:'} {w} × {h} пикс."
            ]

            if hasattr(self, 'ocr_cache'):
                try:
                    cached = self.ocr_cache.get_for_path(path)
                except Exception:
                    cached = None
                if cached:
                    boxes, frames, md5 = cached
                    lines.append(f"{'Кэш:'} Есть (MD5: {md5})")
                    lines.append(f"{' - Боксы:'} {len(boxes)}")
                    lines.append(f"{' - Фреймы:'} {len(frames)}")
                else:
                    lines.append(f"{'Кэш:'} Нет")

            return "\n".join(lines)

        info = get_file_info(path)
        QInputDialog.getMultiLineText(
            self,
            "Информация о файле",
            "Свойства:",
            info
        )

    def refresh_current(self, path=None):
        idx = self.list_widget.currentRow()
        if idx < 0:
            return

        path = Path(self.list_widget.item(idx).data(Qt.UserRole))
        path = str(Path(path).resolve())  

        self.ocr_cache.clear_current(path)

        orig_path = str(self.entries[idx])
        if orig_path != path:
            self.ocr_cache.clear_current(orig_path)

        if hasattr(self, 'current_pixmap') and self.current_pixmap:
            self.ocr_cache.get_for_pixmap(self.current_pixmap)

        self._current_image_token = object()
        self.text_boxes = []
        self.frames = []

        self.show_preview(path)
        self.text_export_panel.set_boxes([], [])
        self.on_item_clicked(self.list_widget.item(idx))

    def action_export_text(self):
        if not hasattr(self, 'ocr_cache') or not self.entries:
            self.statusBar().showMessage("Нет текста для экспорта. Сначала обработайте изображения.")
            return

        first_file = Path(self.entries[0])
        default_name = first_file.parent.name + ".docx"
        filename, _ = QFileDialog.getSaveFileName(self, "Сохранить текст", default_name, "Word Documents (*.docx)")
        if not filename:
            return

        doc = Document()

        for path in self.entries:
            cached = self.ocr_cache.get_for_path(path)
            if cached:
                boxes, frames, _ = cached
                text = "\n".join([box.text for box in boxes if box.text])
                doc.add_paragraph(f"{Path(path).name}")
                doc.add_paragraph(text)
                doc.add_paragraph("") 
        try:
            doc.save(filename)
            self.statusBar().showMessage(f"Текст успешно сохранён в {filename}")
        except Exception as e:
            self.statusBar().showMessage(f"Ошибка сохранения файла: {e}")

    def on_ocr_finished(self, boxes, frames, img_cv, token):
        if token is not None and hasattr(self, '_current_image_token') and token != self._current_image_token:
            return

        self.text_boxes = boxes
        self.frames = frames

        if hasattr(self, 'ocr_cache'):
            try:
                path = self.entries[self.list_widget.currentRow()]
                self.ocr_cache.set_for_path(path, boxes, frames)
            except Exception:
                pass

        self.text_export_panel.set_boxes(boxes, frames=frames)
        self.show_preview(self.entries[self.list_widget.currentRow()], boxes=self.text_export_panel._boxes, reset_zoom=False, frames=frames)

    def show_preview(self, path: Path, boxes=None, frames=None, reset_zoom=False, numbers_boxes=None):
        self.preview_view.text_boxes = boxes

        try:
            pixmap_to_show = QPixmap(str(path))

            if not pixmap_to_show.isNull():
                self.draw_overlays(pixmap_to_show, boxes, frames, numbers_boxes=numbers_boxes)

                if getattr(self, 'current_pixmap_item', None):
                    self.current_pixmap_item.setPixmap(pixmap_to_show)
                else:
                    self.current_pixmap_item = self.scene.addPixmap(pixmap_to_show)

                try:
                    self.scene.setSceneRect(pixmap_to_show.rect())
                except Exception:
                    pass

                if reset_zoom:
                    try:
                        self.preview_view._fit_enabled = True
                        self.preview_view._zoom = 1.0
                    except Exception:
                        pass
                    self.preview_view.fitInView(
                        self.scene.sceneRect(),
                        Qt.KeepAspectRatio
                    )

                self.last_pixmap_for_cache = pixmap_to_show
                self.statusBar().showMessage(str(path))
                try:
                    self.preview_view.update_project_button_visibility()
                except Exception:
                    pass

        except Exception as e:
            self.statusBar().showMessage(f"Ошибка отображения превью: {e}")
            try:
                self.preview_view.update_project_button_visibility()
            except Exception:
                pass

    def draw_overlays(self, pixmap: QPixmap, boxes=None, frames=None, numbers_boxes=None):
        with QPainter(pixmap) as painter:
            painter.setRenderHint(QPainter.Antialiasing, True)
            painter.setRenderHint(QPainter.TextAntialiasing, True)
            if frames and getattr(self, 'show_frames', False):
                pen = QPen(QColor(0, 200, 0, 180), 3)
                painter.setPen(pen)
                painter.setBrush(QColor(0, 200, 0, 40))

                for f in frames:
                    rect = getattr(f, 'rect', f)
                    painter.drawRect(rect)

            numbered_order = {}
            if getattr(self, 'show_numbers_boxes', False):
                panel = getattr(self, 'text_export_panel', None)
                if panel is not None and hasattr(panel, '_boxes') and panel._boxes:
                    for idx, b in enumerate(panel._boxes, start=1):
                        numbered_order[id(b)] = idx
                elif boxes:
                    for idx, b in enumerate(boxes, start=1):
                        numbered_order[id(b)] = idx

            if boxes:
                pen = QPen(QColor(0, 160, 230, 180), 3)
                painter.setPen(pen)
                painter.setBrush(QColor(0, 160, 230, 40))

                metrics = painter.fontMetrics()

                for box in boxes:
                    painter.drawRect(box.rect)

                    if getattr(self, 'show_numbers_boxes', False):
                        idx = numbered_order.get(id(box), None)
                        if idx is not None:
                            num_str = str(idx)
                            saved_font = painter.font()
                            number_font = QFont(saved_font)
                            size_px = max(50, int(min(box.rect.width(), box.rect.height()) * 0.35))
                            number_font.setPixelSize(size_px)
                            number_font.setBold(True)

                            painter.setFont(number_font)
                            fm = painter.fontMetrics()
                            tw = fm.horizontalAdvance(num_str)
                            th = fm.height()
                            center = box.rect.center()

                            path = QPainterPath()
                            path.addText(center.x() - tw / 2, center.y() + th / 4, number_font, num_str)

                            painter.save()
                            painter.setPen(QPen(QColor(255, 255, 255, 220), 10, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
                            painter.setBrush(Qt.NoBrush)
                            painter.drawPath(path)

                            painter.setPen(QPen(QColor(0, 0, 0, 230), 1, Qt.SolidLine, Qt.RoundCap, Qt.RoundJoin))
                            painter.setBrush(QColor(0, 0, 0, 230))
                            painter.drawPath(path)
                            painter.restore()

                            painter.setFont(saved_font)

                    rect = box.rect.adjusted(2, 2, -2, -2)
                    elided_text = metrics.elidedText(
                        box.text, Qt.ElideRight, rect.width()
                    )

                    painter.setFont(QFont())
                    painter.drawText(
                        rect.topLeft() + QPoint(2, metrics.ascent() + 2),
                        elided_text
                    )